#!/usr/bin/env python
"""
The tools file for manipulating msoffice files.

Note, this code was put together using code found within this stackoverflow question:
https://stackoverflow.com/questions/34779724/python-docx-replace-string-in-paragraph-while-keeping-style

"""
from typing import Dict, List


def _docx_paragraph_replace(para_obj, key: str, r_value: str):
    """
    An internal function to replace key values within a docx paragraph
    The reason why you do not replace the text in a paragraph directly
    is that it will cause the original format to change. Replacing the
    text in runs will not cause the original format to change

    :param para_obj: paragraph object from docx.Document object.
    :param key: Keyword that need to be replaced
    :param value: The replaced keywords

    """

    def _r_replace(para_obj, k_maps: List[Dict[str, int]], value: str):
        """
        Accept arguments, removing the characters in k_maps from back to front,
        leaving the first one to replace with value
        Note: Must be removed in reverse order, otherwise the list length
        change will cause IndedxError: string index out of range

        :param k_maps: The list of indexed dictionaries containing keywords，
                       e.g:[{"run":15, "char":3},
                            {"run":15, "char":4},
                            {"run":16, "char":0}]
        :param value: to replace with.

        """
        for i, position in enumerate(reversed(k_maps), start=1):
            y, z = position["run"], position["char"]
            # "k_maps" may contain multiple run ids, which need to be separated
            run: object = para_obj.runs[y]
            # Pit: Instead of the replace() method, str is converted to list
            # after a single word to prevent run.text from making an error in
            # some cases (e.g., a single run contains a duplicate word)
            thisrun = list(run.text)
            if i < len(k_maps):
                # Deleting a corresponding word
                thisrun.pop(z)
            # The last iteration (first word), that is, the number
            # of iterations is equal to the length of k_maps
            if i == len(k_maps):
                # Replace the word in the corresponding position with the new content
                thisrun[z] = value
            run.text = "".join(thisrun)  # Recover

    # Gets the coordinate index values of all the characters
    # in this paragraph [{run_index , char_index}]
    p_maps = [
        {"run": y, "char": z}
        for y, run in enumerate(para_obj.runs)
        for z, char in enumerate(list(run.text))
    ]
    # Handle the number of times key occurs in this paragraph,
    # and record the starting position in the list.
    # Here, while text.find(key) >= 0, the {"ab":"abc"}
    # term will enter an endless loop
    # Takes a single paragraph as an independent body and gets an index
    # list of key positions within the paragraph, or if the paragraph
    # contains multiple keys, there are multiple index values
    k_idx = [
        s
        for s in range(len(para_obj.text))
        if para_obj.text.find(key, s, len(para_obj.text)) == s
    ]

    # Reverse order iteration
    for i, start_idx in enumerate(reversed(k_idx)):
        # The end position of the keyword in this paragraph
        end_idx = start_idx + len(key)
        # Map Slice List A list of dictionaries for
        # sections that contain keywords in a paragraph
        k_maps = p_maps[start_idx:end_idx]
        _r_replace(para_obj, k_maps, r_value)


def docx_replace_body_content(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut body content of a word
    document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    import docx

    docx_obj = docx.Document(word_doc_file)
    print("Replacing within the document body...")
    for key, value in replace_lut.items():
        for para_obj in docx_obj.paragraphs:
            _docx_paragraph_replace(para_obj, key, value)
    docx_obj.save(word_doc_file)


def docx_replace_body_tables(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut tables found within
    the body content of a word document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    import docx

    docx_obj = docx.Document(word_doc_file)

    print("Replacing within the document body tables...")
    for key, value in replace_lut.items():
        for table in docx_obj.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para_obj in cell.paragraphs:
                        _docx_paragraph_replace(para_obj, key, value)
    docx_obj.save(word_doc_file)


def docx_replace_header_content(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut found within
    the header of a word document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    import docx

    docx_obj = docx.Document(word_doc_file)

    print("Replacing within the header body...")
    for key, value in replace_lut.items():
        for section in docx_obj.sections:
            for para_obj in section.header.paragraphs:
                _docx_paragraph_replace(para_obj, key, value)
    docx_obj.save(word_doc_file)


def docx_replace_header_tables(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut found within tables in
    the header of a word document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    import docx

    docx_obj = docx.Document(word_doc_file)

    print("Replacing within the header tables...")
    for key, value in replace_lut.items():
        for section in docx_obj.sections:
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para_obj in cell.paragraphs:
                            _docx_paragraph_replace(para_obj, key, value)
    docx_obj.save(word_doc_file)


def docx_replace_footer_content(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut found within
    the footer of a word document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    import docx

    docx_obj = docx.Document(word_doc_file)

    print("Replacing within the footer body...")
    for key, value in replace_lut.items():
        for section in docx_obj.sections:
            for para_obj in section.footer.paragraphs:
                _docx_paragraph_replace(para_obj, key, value)
    docx_obj.save(word_doc_file)


def docx_replace_footer_tables(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut found within tables in
    the footer of a word document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    import docx

    docx_obj = docx.Document(word_doc_file)

    print("Replacing within footer tables ...")
    for key, value in replace_lut.items():
        for section in docx_obj.sections:
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para_obj in cell.paragraphs:
                            _docx_paragraph_replace(para_obj, key, value)
    docx_obj.save(word_doc_file)


def docx_replace_whole_doc(word_doc_file: str, replace_lut: Dict[str, str]):
    """
    Replace values specified within the replace_lut found within all
    parts (body, header and footer), including tables, within a Word
    document.

    :param word_doc_file: file path to .docx file.
    :param replace_lut: dictionary with replacement values
                        {"value to replace": "with value"}.

    """
    docx_replace_body_content(word_doc_file, replace_lut)
    docx_replace_body_tables(word_doc_file, replace_lut)
    docx_replace_header_content(word_doc_file, replace_lut)
    docx_replace_header_tables(word_doc_file, replace_lut)
    docx_replace_footer_content(word_doc_file, replace_lut)
    docx_replace_footer_tables(word_doc_file, replace_lut)
